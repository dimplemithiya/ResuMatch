import requests
import sys
import json
import os
from datetime import datetime
import subprocess

class ResumeAnalyzerAPITester:
    def __init__(self, base_url="https://skillgap-finder-2.preview.emergentagent.com"):
        self.base_url = base_url
        self.session_token = None
        self.user_id = None
        self.tests_run = 0
        self.tests_passed = 0
        self.analysis_id = None

    def run_test(self, name, method, endpoint, expected_status, data=None, files=None, cookies=None):
        """Run a single API test"""
        url = f"{self.base_url}/api/{endpoint}"
        headers = {'Content-Type': 'application/json'} if not files else {}
        
        if cookies:
            headers.update({'Cookie': f'session_token={cookies}'})

        self.tests_run += 1
        print(f"\n🔍 Testing {name}...")
        print(f"   URL: {url}")
        
        try:
            if method == 'GET':
                response = requests.get(url, headers=headers, cookies={'session_token': cookies} if cookies else None)
            elif method == 'POST':
                if files:
                    response = requests.post(url, data=data, files=files, cookies={'session_token': cookies} if cookies else None)
                else:
                    response = requests.post(url, json=data, headers=headers, cookies={'session_token': cookies} if cookies else None)
            elif method == 'DELETE':
                response = requests.delete(url, headers=headers, cookies={'session_token': cookies} if cookies else None)

            success = response.status_code == expected_status
            if success:
                self.tests_passed += 1
                print(f"✅ Passed - Status: {response.status_code}")
                try:
                    response_data = response.json()
                    print(f"   Response: {json.dumps(response_data, indent=2)[:200]}...")
                    return True, response_data
                except:
                    return True, {}
            else:
                print(f"❌ Failed - Expected {expected_status}, got {response.status_code}")
                try:
                    error_data = response.json()
                    print(f"   Error: {error_data}")
                except:
                    print(f"   Error: {response.text}")
                return False, {}

        except Exception as e:
            print(f"❌ Failed - Error: {str(e)}")
            return False, {}

    def create_test_user_and_session(self):
        """Create test user and session using MongoDB"""
        print("\n🔧 Creating test user and session...")
        
        try:
            # Create test user and session using mongosh
            timestamp = int(datetime.now().timestamp())
            user_id = f"test-user-{timestamp}"
            session_token = f"test_session_{timestamp}"
            email = f"test.user.{timestamp}@example.com"
            
            mongo_script = f"""
use('test_database');
var userId = '{user_id}';
var sessionToken = '{session_token}';
var email = '{email}';
db.users.insertOne({{
  user_id: userId,
  email: email,
  name: 'Test User',
  picture: 'https://via.placeholder.com/150',
  created_at: new Date().toISOString()
}});
db.user_sessions.insertOne({{
  user_id: userId,
  session_token: sessionToken,
  expires_at: new Date(Date.now() + 7*24*60*60*1000).toISOString(),
  created_at: new Date().toISOString()
}});
print('SUCCESS: User and session created');
"""
            
            result = subprocess.run(['mongosh', '--eval', mongo_script], 
                                  capture_output=True, text=True)
            
            if 'SUCCESS' in result.stdout:
                self.session_token = session_token
                self.user_id = user_id
                print(f"✅ Test user created: {user_id}")
                print(f"✅ Session token: {session_token}")
                return True
            else:
                print(f"❌ Failed to create test user: {result.stderr}")
                return False
                
        except Exception as e:
            print(f"❌ Error creating test user: {str(e)}")
            return False

    def test_auth_me(self):
        """Test /api/auth/me endpoint"""
        success, response = self.run_test(
            "Get Current User",
            "GET",
            "auth/me",
            200,
            cookies=self.session_token
        )
        return success

    def test_create_analysis(self):
        """Test resume analysis creation"""
        job_description = """
We are looking for a Senior Software Engineer with:
- 5+ years of experience in React and Node.js
- Strong knowledge of AWS cloud services
- Experience with CI/CD pipelines
- Excellent problem-solving skills
- Knowledge of Docker and Kubernetes
"""
        
        # Create a DOCX file using python-docx
        try:
            from docx import Document
            
            # Create DOCX content
            docx_path = '/tmp/test_resume.docx'
            doc = Document()
            doc.add_heading('John Doe', 0)
            doc.add_heading('Software Engineer', level=1)
            
            doc.add_heading('EXPERIENCE', level=2)
            doc.add_paragraph('• 5 years of React development')
            doc.add_paragraph('• Node.js backend development')
            doc.add_paragraph('• AWS cloud services')
            doc.add_paragraph('• CI/CD pipelines')
            
            doc.add_heading('SKILLS', level=2)
            doc.add_paragraph('• JavaScript, Python')
            doc.add_paragraph('• React, Node.js')
            doc.add_paragraph('• AWS, Docker')
            doc.add_paragraph('• Git, Jenkins')
            
            doc.save(docx_path)
            
            with open(docx_path, 'rb') as f:
                files = {'resume': ('test_resume.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
                data = {'job_description': job_description}
                
                success, response = self.run_test(
                    "Create Analysis",
                    "POST",
                    "analyze",
                    200,
                    data=data,
                    files=files,
                    cookies=self.session_token
                )
                
                if success and 'analysis_id' in response:
                    self.analysis_id = response['analysis_id']
                    print(f"   Analysis ID: {self.analysis_id}")
                
                return success
                
        except ImportError:
            print("   python-docx not available, skipping file upload test")
            print("   ⚠️  File upload functionality needs manual testing")
            return True  # Skip this test but don't fail
                
        except Exception as e:
            print(f"❌ Error in analysis test: {str(e)}")
            return False

    def test_get_analyses(self):
        """Test getting user analyses"""
        success, response = self.run_test(
            "Get User Analyses",
            "GET",
            "analyses",
            200,
            cookies=self.session_token
        )
        return success

    def test_get_analysis_by_id(self):
        """Test getting specific analysis"""
        if not self.analysis_id:
            print("❌ No analysis ID available for testing")
            return False
            
        success, response = self.run_test(
            "Get Analysis by ID",
            "GET",
            f"analyses/{self.analysis_id}",
            200,
            cookies=self.session_token
        )
        return success

    def test_delete_analysis(self):
        """Test deleting analysis"""
        if not self.analysis_id:
            print("❌ No analysis ID available for testing")
            return False
            
        success, response = self.run_test(
            "Delete Analysis",
            "DELETE",
            f"analyses/{self.analysis_id}",
            200,
            cookies=self.session_token
        )
        return success

    def test_logout(self):
        """Test logout endpoint"""
        success, response = self.run_test(
            "Logout",
            "POST",
            "auth/logout",
            200,
            cookies=self.session_token
        )
        return success

    def cleanup_test_data(self):
        """Clean up test data"""
        print("\n🧹 Cleaning up test data...")
        try:
            mongo_script = f"""
use('test_database');
db.users.deleteMany({{email: /test\\.user\\./}});
db.user_sessions.deleteMany({{session_token: /test_session/}});
db.analyses.deleteMany({{user_id: '{self.user_id}'}});
print('SUCCESS: Test data cleaned');
"""
            
            result = subprocess.run(['mongosh', '--eval', mongo_script], 
                                  capture_output=True, text=True)
            
            if 'SUCCESS' in result.stdout:
                print("✅ Test data cleaned successfully")
            else:
                print(f"⚠️  Cleanup warning: {result.stderr}")
                
        except Exception as e:
            print(f"⚠️  Cleanup error: {str(e)}")

def main():
    """Main test execution"""
    print("🚀 Starting AI Resume Analyzer API Tests")
    print("=" * 50)
    
    tester = ResumeAnalyzerAPITester()
    
    # Create test user and session
    if not tester.create_test_user_and_session():
        print("❌ Failed to create test user, stopping tests")
        return 1
    
    # Run API tests
    test_results = []
    
    # Test authentication
    test_results.append(tester.test_auth_me())
    
    # Test analysis workflow
    test_results.append(tester.test_create_analysis())
    test_results.append(tester.test_get_analyses())
    test_results.append(tester.test_get_analysis_by_id())
    
    # Test delete functionality
    test_results.append(tester.test_delete_analysis())
    
    # Test logout
    test_results.append(tester.test_logout())
    
    # Clean up
    tester.cleanup_test_data()
    
    # Print results
    print("\n" + "=" * 50)
    print(f"📊 Tests passed: {tester.tests_passed}/{tester.tests_run}")
    
    if tester.tests_passed == tester.tests_run:
        print("🎉 All tests passed!")
        return 0
    else:
        print("❌ Some tests failed")
        return 1

if __name__ == "__main__":
    sys.exit(main())